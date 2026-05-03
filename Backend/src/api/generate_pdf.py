"""
Generates a Behovsvurdering PDF using the internal RAG answerer.
"""

import io
import re

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_CENTER

from src.api.dependencies import get_current_user
from src.db.database import get_db
from src.db.models import User
from src.rag.answer import get_answerer

router = APIRouter()


# Request schema

class BehovsvurderingFormData(BaseModel):
    typeAnlegg: str = ""
    typeAnleggAnnet: str = ""
    behovBeskrivelse: str = ""
    antallBrukere: str = ""
    kommunalePlanerBeskrivelse: str = ""
    avstandAndreAnlegg: str = ""
    avstandKm: str = ""
    kommune: str = ""
    innbyggertall: str = ""
    befolkningINeromraade: str = ""
    brukerBeskrivelse: str = ""
    idrettslag: str = ""
    antallMedlemmer: str = ""
    driftsansvarlig: str = ""
    driftsmodell: str = ""


# RAG query

def generate_behovsvurdering_text(form: BehovsvurderingFormData) -> str:
    """
    Build a query from the form data and run it through the RAG answerer.
    The answerer uses the local vector store (bestemmelser PDF) as context,
    so the output is grounded in the actual regulations.
    """
    type_anlegg = form.typeAnleggAnnet if form.typeAnlegg == "annet" else form.typeAnlegg

    question = f"""
Skriv en komplett Behovsvurdering for et idrettsanlegg basert på følgende informasjon.
Dokumentet skal dekke alle 8 påkrevde punkter i henhold til gjeldende bestemmelser.
Bruk en formell og saklig tone som passer for offentlige tilskuddssøknader. Svar på norsk bokmål.

Strukturer svaret med nummererte overskrifter slik:
1. Type anlegg
2. Behov (innhold og dimensjonering)
3. Vurdering i kommunens planer
4. Plassering relativt til andre anlegg
5. Innbyggertall og befolkningsdata
6. Beskrivelse av brukerne
7. Idrettslag og medlemstall
8. Driftsansvarlig

Informasjon om anlegget:
- Type anlegg: {type_anlegg}
- Behov og dimensjonering: {form.behovBeskrivelse} | Estimert antall brukere: {form.antallBrukere}
- Kommunale planer: {form.kommunalePlanerBeskrivelse}
- Avstand til andre anlegg: {form.avstandAndreAnlegg} | Nærmeste lignende anlegg: {form.avstandKm} km
- Kommune: {form.kommune} | Innbyggertall: {form.innbyggertall} | Befolkning i nedslagsfelt: {form.befolkningINeromraade}
- Brukerne: {form.brukerBeskrivelse}
- Idrettslag: {form.idrettslag} | Antall medlemmer: {form.antallMedlemmer}
- Driftsansvarlig: {form.driftsansvarlig} | Driftsmodell: {form.driftsmodell}
"""

    answerer = get_answerer()
    result = answerer.answer_with_web(question)  # instead of answerer.answer(question)
    return result["answer"]


# DOCX builder

def build_docx(text: str, kommune: str) -> bytes:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Behovsvurdering")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1B, 0x20, 0x25)

    if kommune:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Kommune: {kommune}").font.size = Pt(11)

    doc.add_paragraph("─" * 60)

    # Parse sections
    sections = re.split(r'(?m)^(\d+\.\s+.+)$', text)

    if len(sections) <= 1:
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        if sections[0].strip():
            doc.add_paragraph(sections[0].strip())
        it = iter(sections[1:])
        for heading in it:
            body = next(it, "")
            h = doc.add_paragraph()
            run = h.add_run(heading.strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1B, 0x20, 0x25)
            for line in body.strip().split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Route

@router.post("/api/generate-pdf")
async def generate_pdf(
    form: BehovsvurderingFormData,
    current_user: User = Depends(get_current_user),
):
    text = generate_behovsvurdering_text(form)
    text = clean_text(text)
    docx_bytes = build_docx(text, form.kommune)
    filename = f"behovsvurdering_{form.kommune or 'dokument'}.docx".replace(" ", "_")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# KOSTNADSOVERSLAG

class KostnadsoverlagFormData(BaseModel):
    prosjektNavn: str = ""
    kommune: str = ""
    typeAnlegg: str = ""
    anleggStorrelse: str = ""
    # Tilskuddsberettigede
    grunnarb: str = ""
    dreneringKr: str = ""
    aktivitetsflateKr: str = ""
    gjerdeUtstyrKr: str = ""
    garderodeKr: str = ""
    lysanleggKr: str = ""
    andreKostnaderTilskudd: str = ""
    # Ikke tilskuddsberettigede
    tribuneKr: str = ""
    parkeringKr: str = ""
    avgifterKr: str = ""
    andreKostnaderIkke: str = ""
    # Dugnad
    dugnadTimer: str = ""
    dugnadTimepris: str = ""
    dugnadBeskrivelse: str = ""
    # Finansiering
    spillemidlerKr: str = ""
    kommunaltTilskuddKr: str = ""
    egneMiddlerKr: str = ""
    andreTilskuddKr: str = ""
    lanKr: str = ""
 
 
def kr(val: str) -> int:
    """Parse a string to int, defaulting to 0."""
    try:
        return int(val or 0)
    except ValueError:
        return 0
 
 
def generate_kostnadsoverlag_text(form: KostnadsoverlagFormData) -> str:
    """Call the RAG answerer to generate a Kostnadsoverslag narrative."""
 
    # Pre-calculate totals to give the model accurate numbers
    tilskudd_total = sum(kr(v) for v in [
        form.grunnarb, form.dreneringKr, form.aktivitetsflateKr,
        form.gjerdeUtstyrKr, form.garderodeKr, form.lysanleggKr,
        form.andreKostnaderTilskudd
    ])
    ikke_tilskudd_total = sum(kr(v) for v in [
        form.tribuneKr, form.parkeringKr, form.avgifterKr, form.andreKostnaderIkke
    ])
    dugnad_total = kr(form.dugnadTimer) * kr(form.dugnadTimepris)
    finansiering_total = sum(kr(v) for v in [
        form.spillemidlerKr, form.kommunaltTilskuddKr,
        form.egneMiddlerKr, form.andreTilskuddKr, form.lanKr
    ])
    prosjekt_total = tilskudd_total + ikke_tilskudd_total + dugnad_total
 
    question = f"""
Skriv et formelt kostnadsoverslag for et idrettsanlegg basert på tallene nedenfor.
Dokumentet skal presentere kostnadene strukturert med disse delene:
1. Prosjektinformasjon
2. Tilskuddsberettigede kostnader (detaljert med poster og totalt)
3. Ikke tilskuddsberettigede kostnader (detaljert med poster og totalt)
4. Dugnad (timer, timepris og totalverdi)
5. Total prosjektkostnad
6. Finansieringsplan (kilder og totalt)
7. Kort vurdering av om prosjektet er fullfinansiert
 
Bruk bestemmelsene i konteksten for å vurdere hvilke kostnader som er tilskuddsberettigede.
Skriv på norsk bokmål med formell og saklig tone. Ingen markdown eller spesialtegn.
 
Prosjektinformasjon:
- Prosjektnavn: {form.prosjektNavn}
- Kommune: {form.kommune}
- Type anlegg: {form.typeAnlegg}
- Storrelse: {form.anleggStorrelse} m2
 
Tilskuddsberettigede kostnader:
- Grunnarbeid og sprengning: {form.grunnarb} kr
- Drenering og avlop: {form.dreneringKr} kr
- Aktivitetsflate/dekke: {form.aktivitetsflateKr} kr
- Gjerde, mal og utstyr: {form.gjerdeUtstyrKr} kr
- Garderobe og sanitar: {form.garderodeKr} kr
- Lysanlegg: {form.lysanleggKr} kr
- Andre tilskuddsberettigede: {form.andreKostnaderTilskudd} kr
- Sum tilskuddsberettigede: {tilskudd_total} kr
 
Ikke tilskuddsberettigede kostnader:
- Tribuneanlegg: {form.tribuneKr} kr
- Parkering og veier: {form.parkeringKr} kr
- Avgifter og gebyrer: {form.avgifterKr} kr
- Andre ikke tilskuddsberettigede: {form.andreKostnaderIkke} kr
- Sum ikke tilskuddsberettigede: {ikke_tilskudd_total} kr
 
Dugnad:
- Antall timer: {form.dugnadTimer}
- Timepris: {form.dugnadTimepris} kr
- Beskrivelse: {form.dugnadBeskrivelse}
- Totalverdi dugnad: {dugnad_total} kr
 
Total prosjektkostnad: {prosjekt_total} kr
 
Finansiering:
- Spillemidler: {form.spillemidlerKr} kr
- Kommunalt tilskudd: {form.kommunaltTilskuddKr} kr
- Egne midler: {form.egneMiddlerKr} kr
- Andre tilskudd: {form.andreTilskuddKr} kr
- Lan: {form.lanKr} kr
- Sum finansiering: {finansiering_total} kr
"""
 
    answerer = get_answerer()
    result = answerer.answer_with_web(question)
    return result["answer"]
 
 
def build_kostnadsoverlag_docx(text: str, form) -> bytes:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Kostnadsoverslag")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1B, 0x20, 0x25)

    subtitle = f"{form.prosjektNavn} - {form.kommune}".strip(" -")
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(subtitle).font.size = Pt(11)

    doc.add_paragraph("─" * 60)

    sections = re.split(r'(?m)^(\d+\.\s+.+)$', text)

    if len(sections) <= 1:
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        if sections[0].strip():
            doc.add_paragraph(sections[0].strip())
        it = iter(sections[1:])
        for heading in it:
            body = next(it, "")
            h = doc.add_paragraph()
            run = h.add_run(heading.strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1B, 0x20, 0x25)
            for line in body.strip().split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def clean_text(text: str) -> str:
    """Replace unsupported Unicode characters with ASCII equivalents."""
    replacements = {
        "•": "-",
        "·": "-",
        "×": "x",
        "–": "-",
        "—": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2022": "-",
        "\u00d7": "x",
        "\u2019": "'",
        "\u2018": "'",
        "\u201c": '"',
        "\u201d": '"',
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text
 
 
@router.post("/api/generate-kostnadsoverlag")
async def generate_kostnadsoverlag(
    form: KostnadsoverlagFormData,
    current_user: User = Depends(get_current_user),
):
    text = generate_kostnadsoverlag_text(form)
    text = clean_text(text)
    docx_bytes = build_kostnadsoverlag_docx(text, form)
    filename = f"kostnadsoverslag_{form.prosjektNavn or form.kommune or 'prosjekt'}.docx".replace(" ", "_")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )